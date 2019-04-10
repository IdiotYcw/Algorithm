# --coding: utf-8
import os
import re
import docx
import csv
import xlrd
from datetime import datetime, date

cwd = os.path.abspath(os.path.curdir)


class Checker(object):
    def __init__(self):
        self.cal_date = date(2018, 12, 31)
        self.mapping = {}
        self.fee_dict = {}
        self.account_list = []
        self.book_dict = {}

    def _find_file(self, keyword):
        for fn in os.listdir(cwd):
            if os.path.isfile(os.path.join(cwd, fn)) and keyword in fn:
                return os.path.abspath(os.path.join(cwd, fn))
        return ''

    def init_mapping(self):
        mapping = self._find_file('基金名称')
        if not mapping:
            raise Exception('NO MAPPING')
        mapping = xlrd.open_workbook(filename=mapping)
        sheets = mapping.sheet_names()
        # print(sheets)
        for sh in sheets:
            self.account_list.append(sh[:2])
            sheet = mapping.sheet_by_name(sh)
            for n in range(1, sheet.nrows):
                row = sheet.row(n)
                if row:
                    name = row[0].value.strip()
                    self.mapping[name] = sh[:2]
        # print(self.mapping)

    def init_fee(self):
        fee_file = self._find_file('费用')
        if not fee_file:
            raise Exception('NO FEE')
        fee = xlrd.open_workbook(filename=fee_file)
        sheet = fee.sheet_by_index(0)
        for n in range(sheet.nrows):
            row = sheet.row(n)
            if row and row[0]:
                if row[0].value and row[0].value.startswith('S'):
                    code = row[0].value.strip()
                    name = row[1].value.strip()[5:]
                    manager_fee = row[2].value
                    sale_fee = row[5].value
                    client_fee = row[6].value
                    self.fee_dict[name] = [code, manager_fee, sale_fee, client_fee]
        # print(self.fee_dict)

    def init_accountbook(self):
        for account in self.account_list:
            print(account)
            file = self._find_file(account)
            if not file:
                continue
                # raise Exception('NO %s ACCOUNTBOOK' % account)
            self.book_dict[account] = []
            book = xlrd.open_workbook(filename=file)
            sheet = book.sheet_by_index(0)

            if '马上' in account:
                for n in range(1, sheet.nrows):
                    row = sheet.row(n)
                    name = row[22].value.strip()
                    scale = row[26].value
                    start_date = row[28].value
                    end_date = row[29].value
                    advance_date = row[30].value
                    start_date = datetime(*xlrd.xldate_as_tuple(
                        start_date, book.datemode)).date() if start_date else None
                    end_date = datetime(*xlrd.xldate_as_tuple(
                        end_date, book.datemode)).date() if end_date else None
                    advance_date = datetime(*xlrd.xldate_as_tuple(
                        advance_date, book.datemode)).date() if advance_date else None
                    self.book_dict[account].append(
                        [name, scale, start_date, end_date, advance_date])
            elif '趣店' in account:
                for n in range(2, sheet.nrows):
                    row = sheet.row(n)
                    name = row[2].value.strip()
                    scale = row[6].value
                    start_date = row[9].value
                    end_date = row[10].value
                    advance_date = row[11].value
                    # print(name, scale)
                    start_date = datetime(*xlrd.xldate_as_tuple(
                        start_date, book.datemode)).date() if start_date else None
                    end_date = datetime(*xlrd.xldate_as_tuple(
                        end_date, book.datemode)).date() if end_date else None
                    advance_date = datetime(*xlrd.xldate_as_tuple(
                        advance_date, book.datemode)).date() if advance_date else None
                    # print(start_date, end_date, advance_date)
                    self.book_dict[account].append(
                        [name, scale, start_date, end_date, advance_date])
            elif '乐信' in account:
                for n in range(2, sheet.nrows):
                    row = sheet.row(n)
                    name = row[2].value.strip()
                    scale = row[6].value
                    start_date = row[9].value
                    end_date = row[10].value
                    advance_date = row[11].value
                    # print(name, scale)
                    start_date = datetime(*xlrd.xldate_as_tuple(
                        start_date, book.datemode)).date() if start_date else None
                    end_date = datetime(*xlrd.xldate_as_tuple(
                        end_date, book.datemode)).date() if end_date else None
                    advance_date = datetime(*xlrd.xldate_as_tuple(
                        advance_date, book.datemode)).date() if advance_date else None
                    # print(start_date, end_date, advance_date)
                    self.book_dict[account].append(
                        [name, scale, start_date, end_date, advance_date])
        # print(self.book_dict)

    def _find_matched_book(self, name):
        return self.mapping.get(name, '')

    def _cal_total_scale(self, table, name):
        scale = 0
        for line in table:
            if line[0] == name and line[2] < self.cal_date < line[3]:
                scale += float(line[1])
        return scale

    # def _find_platform_fee(self, table, code):
    #     fee = {}
    #     for line in table:
    #         if line[0] == code:
    #             return [float(line[2]), float(line[3]), float(line[4])]
    #     return fee

    def _get_nums(self, text):
        text = text.replace('份', '').replace('，', '').replace(',', '')
        return float(text)

    def _get_fees(self, text):
        pattern = r'(?<=【)(.*?)(?=】)'
        sentence = text.split('\n')[-1]
        res = re.findall(pattern, sentence)
        print(res)
        for i, r in enumerate(res):
            res[i] = float(r.replace(',', ''))
        return res

    def start(self):
        for root, _, files in os.walk(os.path.join(cwd, 'doc')):
            if files:
                for f in files:
                    if f.endswith('docx'):
                        p = os.path.join(root, f)
                        print(f)
                        doc = docx.Document(p)

                        name, code, scale = '', '', ''
                        for r, row in enumerate(doc.tables[0].rows):
                            for cell in row.cells:
                                # print(cell.text)
                                if '名称' in cell.text:
                                    name = doc.tables[0].cell(r, 1).text.strip()
                                if '编号' in cell.text:
                                    code = doc.tables[0].cell(r, 1).text.strip()
                                if '份额' in cell.text:
                                    scale = doc.tables[0].cell(r, 1).text.strip()

                        if not name or not code or not scale:
                            print('ERROR!!!')
                            continue

                        book = self._find_matched_book(name)
                        if not book:
                            raise Exception('BOOK ERROR')
                        book_list = self.book_dict.get(book)
                        total_scale = self._cal_total_scale(book_list, name)

                        # print(name, code, scale)
                        doc_scale = self._get_nums(scale)

                        fee_detail = doc.tables[-1].cell(0, 0).text
                        print(fee_detail)
                        fees = self._get_fees(fee_detail)

                        fact_fees = self.fee_dict.get(name, None)
                        if fact_fees is None:
                            raise Exception('FEE ERROR')

                        print("\n\n\n基金： %s\n编号： %s\n\n总份额： %s 份\n管理费： %s元 | 销售服务费： %s元 | 客户服务费： %s元" % (
                            name, code, doc_scale, fees[0], fees[1], fees[2]))

                        print("\n\n\n检测：\n总份额： %s 份\n管理费： %s元 | 销售服务费： %s元 | 客户服务费： %s元" % (
                            total_scale, fact_fees[1], fact_fees[2], fact_fees[3]))

                        scale_check, manager_check, sale_check, client_check = True, True, True, True
                        if doc_scale != total_scale:
                            scale_check = False

                        if fees[0] != fact_fees[1]:
                            manager_check = False
                        if fees[1] != fact_fees[2]:
                            sale_check = False
                        if fees[2] != fact_fees[3]:
                            client_check = False

                        print("\n\n结果:\n\n期末份额： %s\n\n管理费： %s\n\n销售服务费： %s\n\n客户服务费： %s\n\n" % (
                        '一致~' if scale_check else '不一致！！',
                        '一致~' if manager_check else '不一致！！',
                        '一致~' if sale_check else '不一致！！',
                        '一致~' if client_check else '不一致！！'))


if __name__ == '__main__':
    checker = Checker()
    checker.init_mapping()
    checker.init_fee()
    checker.init_accountbook()

    checker.start()
